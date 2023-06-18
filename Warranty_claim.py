import ttkbootstrap as ttk
from tkinter import filedialog
import os , re , csv , math
import python_docx_replace
import docx
import docx2pdf
from docx.shared import Inches
from datetime import datetime
from PIL import Image

class App(ttk.Window):
    def __init__(self):
        super().__init__(
            title="إشعار ضمان",
            themename="flatly",
            size=(700, 780),
            resizable=(True, False),
        )
        # title
        frame = ttk.Frame(self ,bootstyle="dark" )
        frame.pack(fill="both" , anchor="e")
        ttk.Label(frame , text="إشعار ضمان" , font="arial 32 bold" ,  bootstyle="inverse" ,).pack()
        ttk.Separator(self,bootstyle="dark").pack(fill="both" , pady=4)
        # Basic information
        entries = (
            ("Warranty Claim:"      , "warranty_claim"      , "menu"    ,0,0 , ("Accepted - مقبول", "Rejected - مرفوض")),
            ("Distributer:"         , "distribute"          , "entry"   ,1,0 , None),
            ("Location:"            , "location"            , "entry"   ,2,0 , None),
            ("Inspection Center:"   , "inspection_center"   , "entry"   ,3,0 , None),
            ("Technical Inspector:" , "technical_inspector" , "entry"   ,1,1 , None),
            ("Inspection Date:"     , "Inspection_date"     , "date"    ,2,1 , None),
            ("Received On:"         , "received_on"         , "entry"   ,3,1 , None),
        )
        self.basic_info = EntriesFrame(self,"Basic Info (المعلومات الاساسية)",entries)
        # Customer and Vehicle information
        entries = (
            ("Customer Name:"   , "customer_name"   , "entry"   ,0,0 , None),
            ("Customer Address:", "customer_address", "entry"   ,1,0 , None),
            ("Contact Phone:"   , "contact_phone"   , "entry"   ,2,0 , None),
            ("Car:"             , "car"             , "entry"   ,0,1 , None),
            ("Door No:"         , "door_no"         , "entry"   ,1,1 , None),
            ("Contact Email:"   , "contact_email"   , "entry"   ,2,1 , None),
        )
        self.customer_info = EntriesFrame(self,"Customer and Vehicle Info (معلومات العميل و المركبة)",entries)
        # Tire information
        entries = (
            ("Tire:"                , "tire"            , "entry"   ,0,0 , None),
            ("Tire category:"       , "tire_category"   , "entry"   ,1,0 , None),
            ("Load Index [70-170]:" , "load_index"      , "spinbox" ,2,0 , (70,170)),
            ("Origin:"              , "origin"          , "entry"   ,3,0 , None),
            ("Tire Position:"       , "tire_position"   , "menu"    ,1,1 , ("FL","FR","BL","BR")),
            ("Speed Rating:"        , "speed_rating"    , "entry"   ,2,1 , None),
            ("DOT:"                 , "dot"             , "entry"   ,3,1 , None),
        )
        self.tire_info = EntriesFrame(self,"Tire Info (معلومات الإطار)",entries)
        # Damage information
        entries = (
            ("Section:"         , "section"         , "entry"   ,0,0 , None),
            ("Type:"            , "type"            , "entry"   ,1,0 , None),
            ("Detailed Damage:" , "detailed_damage" , "entry"   ,2,0 , None),
            ("PSI:"             , "psi"             , "entry"   ,0,1 , None),
            ("BAR:"             , "bar"             , "entry"   ,1,1 , None),
            ("KPA:"             , "kpa"             , "entry"   ,2,1 , None),
        )
        self.damage_info = EntriesFrame(self,"Damage Info (معلومات الضرر)",entries)
        # add product images
        frame = ttk.Labelframe(self,text="Add Images (إضافة صور)")
        frame.pack(fill="x"  ,padx=2, pady=2)
        headers = ("Image File", "Image Path")
        self.img_table = InfoTable(frame,headers)
        self.img_table.pack(fill="x" , expand=True ,side="left")
        frame = ttk.Frame(frame); frame.pack(side="left",fill="y")
        ttk.Button(frame,text="+" ,bootstyle="outline" , command=self.add_img).pack(fill="both",expand=True)
        ttk.Button(frame,text="-" ,bootstyle="outline" , command=self.img_table.delete_selection).pack(fill="both",expand=True)
        ttk.Button(self,text="PDF اصدار" ,bootstyle="outline" ,command=self.create_pdf_document).pack()
    ###############        ###############        ###############        ###############        
    def add_img(self,):
        files = filedialog.askopenfilenames(filetypes=[('image files', ('.png', '.jpg' , '.jpeg' ))])
        rows = []
        data = tuple(self.img_table.data.values())
        for file in files:
            row = (os.path.basename(file), file)
            if row not in data:
                rows.append(row)
        self.img_table.add_rows(rows)
    ###############        ###############        ###############        ###############    
    def create_pdf_document(self):
        output_file = filedialog.asksaveasfile(filetypes=[('PDF','.pdf')])
        output_file = output_file.name
        if output_file == "":
            return
        if not output_file.lower().endswith('.pdf'):
            output_file += ".pdf"
        data = {}
        with open("data.csv", "r",errors="ignore") as file:
            for row in csv.DictReader(file):
                final_line = row
        row_no = str(int(final_line["id"])+1)
        row_no = row_no.zfill(4)
        additional_info = {"id":row_no,"year":datetime.now().year,"month":datetime.now().month,"day":datetime.now().day}
        data.update(additional_info)
        #get all entries
        data.update(self.basic_info.get_data())
        data.update(self.customer_info.get_data())
        data.update(self.tire_info.get_data())
        data.update(self.damage_info.get_data())
        doc = docx.Document("claims_report_py11.docx")
        python_docx_replace.docx_replace(doc,**data)
        # get images
        img_files = tuple(self.img_table.data.values())
        img_files = img_files[:6]
        TableWithImages(doc,img_files)
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture('stamp.png')
        # create a new document
        temp_file = "temp.docx"
        doc.save(temp_file)
        # convert to pdf
        docx2pdf.convert(temp_file,output_file)
        # Remove word document
        os.remove(temp_file)
        # 
        for key in additional_info.keys():
            if key != "id":
                data.pop(key)
        data["timestamp"] = str(datetime.now())
        with open('data.csv', 'a' , newline="") as f: 
            w = csv.DictWriter(f, data.keys())
            # w.writeheader()
            w.writerow(data)
##############################################################################################################
        
class EntriesFrame(ttk.Labelframe):
    def __init__(self,master,title,entry_ls=()):
        self.entry_dict = {}
        super().__init__(master,text= title,  height=100)
        self.pack(fill="both" , pady =10, padx=2)
        self.entries_frame = ttk.Frame(self); self.entries_frame.pack(fill="both",expand=True,padx=4,pady=4)
        for entry in entry_ls:
            self.add_entry(entry)
    ###############        ###############        ###############        ###############
    def add_entry(self,entry_info):
        label, entry_name , entry_type , row , col , options=entry_info
        self.entries_frame.grid_columnconfigure(col,weight=1)
        frame = ttk.Frame(self.entries_frame)
        frame.grid(sticky="we",row=row,column=col,padx=10)
        ttk.Label(frame , text=f"{label:<20}" , font="arial 10 bold",width=20).pack(side="left" ,anchor="w")
        # entry type
        if entry_type == "entry":
            self.entry_dict[entry_name] = ttk.Entry(frame )
            self.entry_dict[entry_name].pack(side="left", fill="both" , expand=True)
        elif entry_type == "menu":
            self.entry_dict[entry_name] = ttk.StringVar()
            option_menu = ttk.OptionMenu(frame ,self.entry_dict[entry_name] ,options[0], *options , bootstyle="outline")
            option_menu.pack(side="left", fill="both" , expand=True)
            option_menu.config(width=19)
        elif entry_type == "spinbox":
            self.entry_dict[entry_name] = ttk.Spinbox(frame , from_=options[0] , to=options[1])
            self.entry_dict[entry_name].pack(side="left", fill="both" , expand=True)
        elif entry_type == "date":
            DateEntry = ttk.DateEntry(master=frame , dateformat="%d-%m-%Y")
            DateEntry.pack(side="left", fill="both" , expand=True)
            self.entry_dict[entry_name] = DateEntry.entry
    ###############        ###############        ###############        ###############
    def get_data(self):
        data = {}
        for entry_name in self.entry_dict:
            data[entry_name] = self.entry_dict[entry_name].get()
        return data
##############################################################################################################

class InfoTable(ttk.Treeview):
    def __init__(self,master,headers=()):
        self.data = {} # initialize empty tree
        super().__init__(master, columns=headers, show="headings" , bootstyle="primary" , height=3,)
        for header in headers:
            label = header.replace("_", " ")
            label = label.capitalize()
            self.heading(header, text=label)
    ###############        ###############        ###############        ###############
    def clear(self):
        self.data = []
        self.delete(*self.get_children())
    ###############        ###############        ###############        ###############
    def add_rows(self,rows=None):
        if rows is not None:
            for row in rows:
                self.insert('', ttk.END, values=row)
                self.data[self.get_children()[-1]] = row
    ###############        ###############        ###############        ###############
    def add_new_rows(self,rows=None):
        if rows is not None:
            self.data = {}
            self.clear()
            self.add_rows(rows)
    ###############        ###############        ###############        ###############
    def delete_selection(self):
        for sel_item in self.selection():
            self.delete(sel_item)
            self.data.pop(sel_item)
##############################################################################################################

class NewRow(ttk.Frame):
    def __init__(self,master,column_percentage=(),padx=0,pady=0):
        super().__init__(master)
        self.pack(fill="both",padx=padx,pady=pady)
        for column_no,column_wight in enumerate(column_percentage):
            self.columnconfigure(column_no,weight=column_wight)
##############################################################################################################

class TableWithImages():
    def __init__(self,doc,imgs_path):
        for img_no , img in enumerate(imgs_path):
            _ , file = img
            cell = (img_no)%2
            if cell == 0:
                table = doc.add_table(rows=1, cols=2 , style='Table Grid')
            p = table.rows[0].cells[cell].add_paragraph()
            r = p.add_run()
            w_inch , h_inch = self.resize_img_inch(file,250,250)
            r.add_picture(file , width=Inches(w_inch), height=Inches(h_inch))
    ###############        ###############        ###############        ###############
    def resize_img(self,file_path,max_w,max_h):
        img = Image.open(file_path)
        w,h =img.size
        if w>h:
            h= h *max_w/w
            w=max_w
        else : 
            w = w*max_h/h
            h = max_h
        return w,h
    ###############        ###############        ###############        ###############
    def resize_img_inch(self,file_path,max_w,max_h):
        w,h = self.resize_img(file_path,max_w,max_h)
        return w/96 , h/96 
##############################################################################################################

if __name__ == '__main__':
    app = App()
    app.mainloop()
